from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE = Path(__file__).with_name("course.txt")
TEMPLATE = Path(__file__).with_name("Хоронеко_Дмитрий_1_группа.docx")
TARGET = Path(__file__).with_name("КУРСОВАЯ_РАБОТА_ГОСТ.docx")

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
SEPARATOR = "=" * 80


def clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_run_font(run, *, bold: bool = False, size: Pt = FONT_SIZE) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_text_run(paragraph, text: str, *, bold: bool = False, size: Pt = FONT_SIZE):
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return run


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    heading1 = document.styles["Heading 1"]
    heading1.font.name = FONT_NAME
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(0)
    heading1.paragraph_format.line_spacing = 1.0

    heading2 = document.styles["Heading 2"]
    heading2.font.name = FONT_NAME
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    heading2.paragraph_format.space_before = Pt(2)
    heading2.paragraph_format.space_after = Pt(0)
    heading2.paragraph_format.line_spacing = 1.0

    title = document.styles["Title"]
    title.font.name = FONT_NAME
    title.font.size = Pt(16)
    title.font.bold = True
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_title_page(document: Document, lines: list[str]) -> None:
    cleaned = [line.rstrip() for line in lines]
    while cleaned and not cleaned[-1]:
        cleaned.pop()

    for line in cleaned:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        if line:
            if line == "ВЕБ-КАТАЛОГ ЗООМАГАЗИНА":
                add_text_run(paragraph, line, bold=True, size=Pt(16))
            else:
                add_text_run(paragraph, line)

    document.add_page_break()


def add_toc_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run_begin = paragraph.add_run()
    run_begin._r.append(begin)

    run_instr = paragraph.add_run()
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    run_sep._r.append(separate)

    placeholder = paragraph.add_run("Обновите оглавление в Word: выделите таблицу и нажмите F9.")
    set_run_font(placeholder)

    run_end = paragraph.add_run()
    run_end._r.append(end)


def add_table_of_contents(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Cm(0)
    add_text_run(heading, "ОГЛАВЛЕНИЕ", bold=True, size=Pt(16))

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    add_toc_field(paragraph)
    document.add_page_break()


def is_heading1(line: str) -> bool:
    return (
        line == "ВВЕДЕНИЕ"
        or line == "ЗАКЛЮЧЕНИЕ"
        or line == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
        or line.startswith("ГЛАВА ")
    )


def is_heading2(line: str) -> bool:
    return re.match(r"^\d+\.\d+\s+", line) is not None


def add_heading1(document: Document, text: str, *, first_body_heading: bool) -> None:
    if not first_body_heading:
        document.add_page_break()
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_text_run(paragraph, text, bold=True, size=Pt(16))


def add_heading2(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(1.25)
    add_text_run(paragraph, text, bold=True)


def add_body_line(document: Document, text: str) -> None:
    is_list = text.startswith("– ") or re.match(r"^\d+\.\s", text) is not None
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(0) if is_list else Cm(1.25)
    add_text_run(paragraph, text)


def split_source(lines: list[str]) -> tuple[list[str], list[str]]:
    first_sep = lines.index(SEPARATOR)
    second_sep = lines.index(SEPARATOR, first_sep + 1)
    title_lines = lines[:first_sep]
    body_lines = lines[second_sep + 1 :]
    return title_lines, body_lines


def build_document(source: Path, target: Path) -> None:
    document = Document(str(TEMPLATE))
    clear_document_body(document)
    configure_document(document)
    enable_update_fields_on_open(document)

    lines = source.read_text(encoding="utf-8").splitlines()
    title_lines, body_lines = split_source(lines)

    add_title_page(document, title_lines)
    add_table_of_contents(document)

    first_body_heading = True
    for raw_line in body_lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped or stripped == SEPARATOR:
            continue

        if is_heading1(stripped):
            add_heading1(document, stripped, first_body_heading=first_body_heading)
            first_body_heading = False
            continue

        if is_heading2(stripped):
            add_heading2(document, stripped)
            continue

        add_body_line(document, stripped)

    document.save(target)


if __name__ == "__main__":
    build_document(SOURCE, TARGET)
    print(f"Created: {TARGET}")
